import streamlit as st
import gspread
import requests 
from oauth2client.service_account import ServiceAccountCredentials
import folium
from streamlit_folium import folium_static
import pandas as pd
from docx import Document
from io import BytesIO
import zipfile
from streamlit_option_menu import option_menu
from streamlit_gsheets import GsheetsConnection
from PIL import Image
from io import BytesIO
import plotly.graph_objects as go
import plotly.express as px
import os
import json

#=====================================================================================================
# Nama spreadsheet
spreadsheet_name = 'SITACSULAWESI'

# Ambil konfigurasi dari file TOML
config = st.experimental_get_config()
gsheets_config = config["gsheets"]

# Fungsi untuk menghubungkan ke Google Sheets
def connect_to_google_sheets(spreadsheet_name, sheet_name):
    scope = [
        'https://spreadsheets.google.com/feeds',
        'https://www.googleapis.com/auth/drive'
    ]

    # Load credentials from TOML configuration
    creds_dict = gsheets_config  # Already loaded from TOML

    # Verify the credentials dictionary (optional, as TOML should be correctly formatted)
    required_keys = [
        "type", "project_id", "private_key_id", "private_key", "client_email",
        "client_id", "auth_uri", "token_uri", "auth_provider_x509_cert_url",
        "client_x509_cert_url"
    ]
    for key in required_keys:
        if key not in creds_dict:
            raise ValueError(f"Key '{key}' is missing from the credentials TOML")

    creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
    client = gspread.authorize(creds)

    # Buka spreadsheet berdasarkan nama
    spreadsheet = client.open(spreadsheet_name)

    # Pilih worksheet berdasarkan nama
    try:
        worksheet = spreadsheet.worksheet(sheet_name)
    except gspread.exceptions.WorksheetNotFound:
        raise ValueError(f"Worksheet '{sheet_name}' tidak ditemukan di spreadsheet '{spreadsheet_name}'")

    return worksheet
#=====================================================================================================

# Fungsi untuk menampilkan data berdasarkan Site ID Operator
def show_tssr_data(sheet, site_id_operator):
    st.title("Show TSSR")

    # Mendapatkan semua data dari sheet
    data = sheet.get_all_values()

    # Menampilkan data berdasarkan Site ID Operator
    headers = data[1]
    site_id_col_index = headers.index("Site ID Operator")
    sonumb_col_index = headers.index("Sonumb")
    site_name_col_index = headers.index("Site Name TBG")
    long_nom_col_index = headers.index("Long NOM")
    lat_nom_col_index = headers.index("Lat NOM")
    mitra_col_index = headers.index("Mitra")
    
    # Indexes for "Kandidat P"
    alamat_kand_p_col_index = headers.index("Alamat Site Kand P")
    long_kand_p_col_index = headers.index("Long Kand P")
    lat_kand_p_col_index = headers.index("Lat Kand P")
    jarak_kand_p_col_index = headers.index("Jarak kand P dari titik NOM (meter)")
    Kepemilikan_Lahan_Kand_P_index = headers.index("Kepemilikan Lahan Kand P")
    Luas_Lahan_Kand_P_meter_index	= headers.index("Luas Lahan Kand P (meter)")
    Nama_Pemilik_Lahan_Kand_P_index	= headers.index("Nama Pemilik Lahan Kand P")
    Alamat_Pemilik_Lahan_Kand_P_index = headers.index("Alamat Pemilik Lahan Kand P")
    POIKawasan_ramai_disekitar_Kand_P_index = headers.index("POI/Kawasan ramai disekitar Kand P")
    Tipe_Area_sekitar_Kand_P_index = headers.index("Tipe Area sekitar Kand P")
    Tipe_Site_Kand_P_index = headers.index("Tipe Site Kand P")
    Akses_Menuju_Kand_P_index = headers.index("Akses Menuju Kand P")
    ukuran_akses_ke_Kand_P_meter_index = headers.index("ukuran akses ke Kand P (meter)")
    Site_terdekat_dengan_kandidat_P_operator_index = headers.index("Site terdekat dengan kandidat P (operator)")
    Jarak_site_terdekat_dengan_kandidat_P_meter_index	= headers.index("Jarak site terdekat dengan kandidat P (meter)")
    Tipe_site_terdekat_dengan_kandidat_P_index = headers.index("Tipe site terdekat dengan kandidat P")
    Tipe_Tower_terdekat_Kand_P_index = headers.index("Tipe Tower terdekat Kand P")
    Long_site_terdekat_kand_P_index	= headers.index("Long site terdekat kand P")
    Lat_site_terdekat_kand_P_index	= headers.index("Lat site terdekat kand P")
    Informasi_lain_Kand_P_index	= headers.index("Informasi lain Kand P")
    Jarak_terdekat_tiang_PLN_ke_Kand_P_meter_index = headers.index("Jarak terdekat tiang PLN ke Kand P (meter)")
    Jarak_terdekat_transformator_ke_kand_P_meter_index = headers.index("Jarak terdekat transformator ke kand P (meter)")
    Validation_Kand_P_index = headers.index("Validation Kand P")
   	
    
    # Indexes for "Kandidat Q"
    alamat_kand_q_col_index = headers.index("Alamat Site Kand Q")
    long_kand_q_col_index = headers.index("Long Kand Q")
    lat_kand_q_col_index = headers.index("Lat Kand Q")
    jarak_kand_q_col_index = headers.index("Jarak kand Q dari titik NOM (meter)")
    Kepemilikan_Lahan_Kand_Q_index = headers.index("Kepemilikan Lahan Kand Q")
    Luas_Lahan_Kand_Q_meter_index	= headers.index("Luas Lahan Kand Q (meter)")
    Nama_Pemilik_Lahan_Kand_Q_index	= headers.index("Nama Pemilik Lahan Kand Q")
    Alamat_Pemilik_Lahan_Kand_Q_index = headers.index("Alamat Pemilik Lahan Kand Q")
    POIKawasan_ramai_disekitar_Kand_Q_index = headers.index("POI/Kawasan ramai disekitar Kand Q")
    Tipe_Area_sekitar_Kand_Q_index = headers.index("Tipe Area sekitar Kand Q")
    Tipe_Site_Kand_Q_index = headers.index("Tipe Site Kand Q")
    Akses_Menuju_Kand_Q_index = headers.index("Akses Menuju Kand Q")
    ukuran_akses_ke_Kand_Q_meter_index = headers.index("ukuran akses ke Kand Q (meter)")
    Site_terdekat_dengan_kandidat_Q_operator_index = headers.index("Site terdekat dengan kandidat Q (operator)")
    Jarak_site_terdekat_dengan_kandidat_Q_meter_index	= headers.index("Jarak site terdekat dengan kandidat Q (meter)")
    Tipe_site_terdekat_dengan_kandidat_Q_index = headers.index("Tipe site terdekat dengan kandidat Q")
    Tipe_Tower_terdekat_Kand_Q_index = headers.index("Tipe Tower terdekat Kand Q")
    Long_site_terdekat_kand_Q_index	= headers.index("Long site terdekat kand Q")
    Lat_site_terdekat_kand_Q_index	= headers.index("Lat site terdekat kand Q")
    Informasi_lain_Kand_Q_index	= headers.index("Informasi lain Kand Q")
    Jarak_terdekat_tiang_PLN_ke_Kand_Q_meter_index = headers.index("Jarak terdekat tiang PLN ke Kand Q (meter)")
    Jarak_terdekat_transformator_ke_kand_Q_meter_index = headers.index("Jarak terdekat transformator ke kand Q (meter)")
    Validation_Kand_Q_index = headers.index("Validation Kand Q")
    

    # Indexes for "Kandidat R"
    alamat_kand_r_col_index = headers.index("Alamat Site Kand R")
    long_kand_r_col_index = headers.index("Long Kand R")
    lat_kand_r_col_index = headers.index("Lat Kand R")
    jarak_kand_r_col_index = headers.index("Jarak kand R dari titik NOM (meter)")
    Kepemilikan_Lahan_Kand_r_index = headers.index("Kepemilikan Lahan Kand R")
    Luas_Lahan_Kand_r_meter_index	= headers.index("Luas Lahan Kand R (meter)")
    Nama_Pemilik_Lahan_Kand_r_index	= headers.index("Nama Pemilik Lahan Kand R")
    Alamat_Pemilik_Lahan_Kand_r_index = headers.index("Alamat Pemilik Lahan Kand R")
    POIKawasan_ramai_disekitar_Kand_r_index = headers.index("POI/Kawasan ramai disekitar Kand R")
    Tipe_Area_sekitar_Kand_r_index = headers.index("Tipe Area sekitar Kand R")
    Tipe_Site_Kand_r_index = headers.index("Tipe Site Kand R")
    Akses_Menuju_Kand_r_index = headers.index("Akses Menuju Kand R")
    ukuran_akses_ke_Kand_r_meter_index = headers.index("ukuran akses ke Kand R (meter)")
    Site_terdekat_dengan_kandidat_r_operator_index = headers.index("Site terdekat dengan kandidat R (operator)")
    Jarak_site_terdekat_dengan_kandidat_r_meter_index = headers.index("Jarak site terdekat dengan kandidat R (meter)")
    Tipe_site_terdekat_dengan_kandidat_r_index = headers.index("Tipe site terdekat dengan kandidat R")
    Tipe_Tower_terdekat_Kand_r_index = headers.index("Tipe Tower terdekat Kand R")
    Long_site_terdekat_kand_r_index	= headers.index("Long site terdekat kand R")
    Lat_site_terdekat_kand_r_index	= headers.index("Lat site terdekat kand R")
    Informasi_lain_Kand_r_index	= headers.index("Informasi lain Kand R")
    Jarak_terdekat_tiang_PLN_ke_Kand_r_meter_index = headers.index("Jarak terdekat tiang PLN ke Kand R (meter)")
    Jarak_terdekat_transformator_ke_kand_r_meter_index = headers.index("Jarak terdekat transformator ke kand R (meter)")
    Validation_Kand_r_index = headers.index("Validation Kand R")
    
    # Indexes for "Kandidat S"
    alamat_kand_s_col_index = headers.index("Alamat Site Kand S")
    long_kand_s_col_index = headers.index("Long Kand S")
    lat_kand_s_col_index = headers.index("Lat Kand S")
    jarak_kand_s_col_index = headers.index("Jarak kand S dari titik NOM (meter)")
    Kepemilikan_Lahan_Kand_s_index = headers.index("Kepemilikan Lahan Kand S")
    Luas_Lahan_Kand_s_meter_index	= headers.index("Luas Lahan Kand S (meter)")
    Nama_Pemilik_Lahan_Kand_s_index	= headers.index("Nama Pemilik Lahan Kand S")
    Alamat_Pemilik_Lahan_Kand_s_index = headers.index("Alamat Pemilik Lahan Kand S")
    POIKawasan_ramai_disekitar_Kand_s_index = headers.index("POI/Kawasan ramai disekitar Kand S")
    Tipe_Area_sekitar_Kand_s_index = headers.index("Tipe Area sekitar Kand S")
    Tipe_Site_Kand_s_index = headers.index("Tipe Site Kand S")
    Akses_Menuju_Kand_s_index = headers.index("Akses Menuju Kand S")
    ukuran_akses_ke_Kand_s_meter_index = headers.index("ukuran akses ke Kand S (meter)")
    Site_terdekat_dengan_kandidat_s_operator_index = headers.index("Site terdekat dengan kandidat S (operator)")
    Jarak_site_terdekat_dengan_kandidat_s_meter_index = headers.index("Jarak site terdekat dengan kandidat S (meter)")
    Tipe_site_terdekat_dengan_kandidat_s_index = headers.index("Tipe site terdekat dengan kandidat S")
    Tipe_Tower_terdekat_Kand_s_index = headers.index("Tipe Tower terdekat Kand S")
    Long_site_terdekat_kand_s_index	= headers.index("Long site terdekat kand S")
    Lat_site_terdekat_kand_s_index	= headers.index("Lat site terdekat kand S")
    Informasi_lain_Kand_s_index	= headers.index("Informasi lain Kand S")
    Jarak_terdekat_tiang_PLN_ke_Kand_s_meter_index = headers.index("Jarak terdekat tiang PLN ke Kand S (meter)")
    Jarak_terdekat_transformator_ke_kand_s_meter_index = headers.index("Jarak terdekat transformator ke kand S (meter)")
    Validation_Kand_s_index = headers.index("Validation Kand S")

    for row in data[2:]:
        if row[site_id_col_index] == site_id_operator:
            st.session_state["tssr_data"] = {
                "Sonumb": row[sonumb_col_index],
                "Site Name TBG": row[site_name_col_index],
                "Long NOM": row[long_nom_col_index],
                "Lat NOM": row[lat_nom_col_index],
                "Mitra": row[mitra_col_index],
                "Alamat Site Kand P": row[alamat_kand_p_col_index],
                "Long Kand P": row[long_kand_p_col_index],
                "Lat Kand P": row[lat_kand_p_col_index],
                "Jarak kand P dari titik NOM (meter)": row[jarak_kand_p_col_index],
                "Kepemilikan Lahan Kand P": row[Kepemilikan_Lahan_Kand_P_index],
                "Luas Lahan Kand P (meter)": row[Luas_Lahan_Kand_P_meter_index],
                "Nama Pemilik Lahan Kand P": row[Nama_Pemilik_Lahan_Kand_P_index],
                "Alamat Pemilik Lahan Kand P": row[Alamat_Pemilik_Lahan_Kand_P_index],
                "POI/Kawasan ramai disekitar Kand P": row[POIKawasan_ramai_disekitar_Kand_P_index],
                "Tipe Area sekitar Kand P": row[Tipe_Area_sekitar_Kand_P_index],
                "Tipe Site Kand P": row[Tipe_Site_Kand_P_index],
                "Akses Menuju Kand P": row[Akses_Menuju_Kand_P_index],
                "ukuran akses ke Kand P (meter)": row[ukuran_akses_ke_Kand_P_meter_index],
                "Site terdekat dengan kandidat P (operator)": row[Site_terdekat_dengan_kandidat_P_operator_index],
                "Jarak site terdekat dengan kandidat P (meter)": row[Jarak_site_terdekat_dengan_kandidat_P_meter_index],
                "Tipe site terdekat dengan kandidat P": row[Tipe_site_terdekat_dengan_kandidat_P_index],
                "Tipe Tower terdekat Kand P": row[Tipe_Tower_terdekat_Kand_P_index],
                "Long site terdekat kand P": row[Long_site_terdekat_kand_P_index],
                "Lat site terdekat kand P": row[Lat_site_terdekat_kand_P_index],
                "Informasi lain Kand P": row[Informasi_lain_Kand_P_index],
                "Jarak terdekat tiang PLN ke Kand P (meter)": row[Jarak_terdekat_tiang_PLN_ke_Kand_P_meter_index],
                "Jarak terdekat transformator ke kand P (meter)": row[Jarak_terdekat_transformator_ke_kand_P_meter_index],
                "Validation Kand P": row[Validation_Kand_P_index],
                "Alamat Site Kand Q": row[alamat_kand_q_col_index],
                "Long Kand Q": row[long_kand_q_col_index],
                "Lat Kand Q": row[lat_kand_q_col_index],
                "Jarak kand Q dari titik NOM (meter)": row[jarak_kand_q_col_index],
                "Kepemilikan Lahan Kand Q": row[Kepemilikan_Lahan_Kand_Q_index],
                "Luas Lahan Kand Q (meter)": row[Luas_Lahan_Kand_Q_meter_index],
                "Nama Pemilik Lahan Kand Q": row[Nama_Pemilik_Lahan_Kand_Q_index],
                "Alamat Pemilik Lahan Kand Q": row[Alamat_Pemilik_Lahan_Kand_Q_index],
                "POI/Kawasan ramai disekitar Kand Q": row[POIKawasan_ramai_disekitar_Kand_Q_index],
                "Tipe Area sekitar Kand Q": row[Tipe_Area_sekitar_Kand_Q_index],
                "Tipe Site Kand Q": row[Tipe_Site_Kand_Q_index],
                "Akses Menuju Kand Q": row[Akses_Menuju_Kand_Q_index],
                "ukuran akses ke Kand Q (meter)": row[ukuran_akses_ke_Kand_Q_meter_index],
                "Site terdekat dengan kandidat Q (operator)": row[Site_terdekat_dengan_kandidat_Q_operator_index],
                "Jarak site terdekat dengan kandidat Q (meter)": row[Jarak_site_terdekat_dengan_kandidat_Q_meter_index],
                "Tipe site terdekat dengan kandidat Q": row[Tipe_site_terdekat_dengan_kandidat_Q_index],
                "Tipe Tower terdekat Kand Q": row[Tipe_Tower_terdekat_Kand_Q_index],
                "Long site terdekat kand Q": row[Long_site_terdekat_kand_Q_index],
                "Lat site terdekat kand Q": row[Lat_site_terdekat_kand_Q_index],
                "Informasi lain Kand Q": row[Informasi_lain_Kand_Q_index],
                "Jarak terdekat tiang PLN ke Kand Q (meter)": row[Jarak_terdekat_tiang_PLN_ke_Kand_Q_meter_index],
                "Jarak terdekat transformator ke kand Q (meter)": row[Jarak_terdekat_transformator_ke_kand_Q_meter_index],
                "Validation Kand Q": row[Validation_Kand_Q_index],
                "Alamat Site Kand R": row[alamat_kand_r_col_index],
                "Long Kand R": row[long_kand_r_col_index],
                "Lat Kand R": row[lat_kand_r_col_index],
                "Jarak kand R dari titik NOM (meter)": row[jarak_kand_r_col_index],
                "Kepemilikan Lahan Kand R": row[Kepemilikan_Lahan_Kand_r_index],
                "Luas Lahan Kand R (meter)": row[Luas_Lahan_Kand_r_meter_index],
                "Nama Pemilik Lahan Kand R": row[Nama_Pemilik_Lahan_Kand_r_index],
                "Alamat Pemilik Lahan Kand R": row[Alamat_Pemilik_Lahan_Kand_r_index],
                "POI/Kawasan ramai disekitar Kand R": row[POIKawasan_ramai_disekitar_Kand_r_index],
                "Tipe Area sekitar Kand R": row[Tipe_Area_sekitar_Kand_r_index],
                "Tipe Site Kand R": row[Tipe_Site_Kand_r_index],
                "Akses Menuju Kand R": row[Akses_Menuju_Kand_r_index],
                "ukuran akses ke Kand R (meter)": row[ukuran_akses_ke_Kand_r_meter_index],
                "Site terdekat dengan kandidat R (operator)": row[Site_terdekat_dengan_kandidat_r_operator_index],
                "Jarak site terdekat dengan kandidat R (meter)": row[Jarak_site_terdekat_dengan_kandidat_r_meter_index],
                "Tipe site terdekat dengan kandidat R": row[Tipe_site_terdekat_dengan_kandidat_r_index],
                "Tipe Tower terdekat Kand R": row[Tipe_Tower_terdekat_Kand_r_index],
                "Long site terdekat kand R": row[Long_site_terdekat_kand_r_index],
                "Lat site terdekat kand R": row[Lat_site_terdekat_kand_r_index],
                "Informasi lain Kand R": row[Informasi_lain_Kand_r_index],
                "Jarak terdekat tiang PLN ke Kand R (meter)": row[Jarak_terdekat_tiang_PLN_ke_Kand_r_meter_index],
                "Jarak terdekat transformator ke kand R (meter)": row[Jarak_terdekat_transformator_ke_kand_r_meter_index],
                "Validation Kand R": row[Validation_Kand_r_index],
                "Alamat Site Kand S": row[alamat_kand_s_col_index],
                "Long Kand S": row[long_kand_s_col_index],
                "Lat Kand S": row[lat_kand_s_col_index],
                "Jarak kand S dari titik NOM (meter)": row[jarak_kand_s_col_index],
                "Kepemilikan Lahan Kand S": row[Kepemilikan_Lahan_Kand_s_index],
                "Luas Lahan Kand S (meter)": row[Luas_Lahan_Kand_s_meter_index],
                "Nama Pemilik Lahan Kand S": row[Nama_Pemilik_Lahan_Kand_s_index],
                "Alamat Pemilik Lahan Kand S": row[Alamat_Pemilik_Lahan_Kand_s_index],
                "POI/Kawasan ramai disekitar Kand S": row[POIKawasan_ramai_disekitar_Kand_s_index],
                "Tipe Area sekitar Kand S": row[Tipe_Area_sekitar_Kand_s_index],
                "Tipe Site Kand S": row[Tipe_Site_Kand_s_index],
                "Akses Menuju Kand S": row[Akses_Menuju_Kand_s_index],
                "ukuran akses ke Kand S (meter)": row[ukuran_akses_ke_Kand_s_meter_index],
                "Site terdekat dengan kandidat S (operator)": row[Site_terdekat_dengan_kandidat_s_operator_index],
                "Jarak site terdekat dengan kandidat S (meter)": row[Jarak_site_terdekat_dengan_kandidat_s_meter_index],
                "Tipe site terdekat dengan kandidat S": row[Tipe_site_terdekat_dengan_kandidat_s_index],
                "Tipe Tower terdekat Kand S": row[Tipe_Tower_terdekat_Kand_s_index],
                "Long site terdekat kand S": row[Long_site_terdekat_kand_s_index],
                "Lat site terdekat kand S": row[Lat_site_terdekat_kand_s_index],
                "Informasi lain Kand S": row[Informasi_lain_Kand_s_index],
                "Jarak terdekat tiang PLN ke Kand S (meter)": row[Jarak_terdekat_tiang_PLN_ke_Kand_s_meter_index],
                "Jarak terdekat transformator ke kand S (meter)": row[Jarak_terdekat_transformator_ke_kand_s_meter_index],
                "Validation Kand S": row[Validation_Kand_s_index],
            }
            break
    else:
        st.session_state["tssr_data"] = None

# Fungsi untuk menampilkan peta
def show_map(lat, long, lat_kand_p=None, long_kand_p=None, lat_kand_q=None, long_kand_q=None, lat_kand_r=None, long_kand_r=None, lat_kand_s=None, long_kand_s=None, radius=None):
    m = folium.Map(location=[lat, long], zoom_start=15)
    
    # Penanda titik NOM
    folium.Marker([lat, long], popup="NOM", tooltip="NOM").add_to(m)
    
    # Penanda titik Kandidat P dengan huruf "P" jika koordinat ada
    if lat_kand_p is not None and long_kand_p is not None:
        folium.Marker(
            [lat_kand_p, long_kand_p],
            popup="Kandidat P",
            tooltip="Kandidat P",
            icon=folium.Icon(color="red", icon='P', prefix='fa')
        ).add_to(m)

    # Penanda titik Kandidat Q dengan huruf "Q" jika koordinat ada
    if lat_kand_q is not None and long_kand_q is not None:
        folium.Marker(
            [lat_kand_q, long_kand_q],
            popup="Kandidat Q",
            tooltip="Kandidat Q",
            icon=folium.Icon(color="green", icon='Q', prefix='fa')
        ).add_to(m)

    # Penanda titik Kandidat R dengan huruf "R" jika koordinat ada
    if lat_kand_r is not None and long_kand_r is not None:
        folium.Marker(
            [lat_kand_r, long_kand_r],
            popup="Kandidat R",
            tooltip="Kandidat R",
            icon=folium.Icon(color="blue", icon='R', prefix='fa')
        ).add_to(m)

    # Penanda titik Kandidat S dengan huruf "S" jika koordinat ada
    if lat_kand_s is not None and long_kand_s is not None:
        folium.Marker(
            [lat_kand_s, long_kand_s],
            popup="Kandidat S",
            tooltip="Kandidat S",
            icon=folium.Icon(color="purple", icon='S', prefix='fa')
        ).add_to(m)
    
    if radius:
        folium.Circle(
            location=[lat, long],
            radius=radius,
            color='blue',
            fill=True,
            fill_color='blue'
        ).add_to(m)
    
    folium_static(m)
    
    

# Fungsi untuk menampilkan dashboard progress monitoring
def show_progress_dashboard(sheet):
    st.title("Progress Monitoring Dashboard")

    # Mendapatkan semua data dari sheet
    data = sheet.get_all_values()

    # Mengubah data menjadi DataFrame untuk kemudahan pengolahan
    df = pd.DataFrame(data[1:], columns=data[0])

    # Menghitung jumlah total Site ID Operator
    total_sites = df["Site ID Operator"].nunique()

    # Menghitung jumlah Site ID Operator berdasarkan Mitra
    site_counts_by_mitra = df.groupby("Mitra")["Site ID Operator"].nunique().reset_index()

    # Menghitung jumlah Candidate NY Submit by Mitra
    candidate_submit_by_mitra = df[df["LT Candidate Submitted NY Validation"] == "Candidate NY Submit by Mitra"].groupby("Mitra").size().reset_index(name="Candidate NY Submit")

    # Menggabungkan site_counts_by_mitra dengan candidate_submit_by_mitra
    site_counts_by_mitra = site_counts_by_mitra.merge(candidate_submit_by_mitra, on="Mitra", how="left").fillna(0)

    # Menghitung Need Validation By Operator dan persentase
    site_counts_by_mitra["Need Validation By Operator"] = site_counts_by_mitra["Site ID Operator"] - site_counts_by_mitra["Candidate NY Submit"]
    site_counts_by_mitra["Percentage"] = site_counts_by_mitra.apply(lambda row: (row["Site ID Operator"] - row["Candidate NY Submit"]) / row["Site ID Operator"] * 100, axis=1)

    # Menampilkan hasil di dashboard
    st.write(f"Total Site ID Operator: {total_sites}")

    # Plot grafik batang untuk setiap Mitra
    fig = px.bar(site_counts_by_mitra, x="Mitra", y=["Site ID Operator", "Candidate NY Submit", "Need Validation By Operator"],
                 barmode='group', title="Progress Monitoring per Mitra")
    st.plotly_chart(fig)

    for index, row in site_counts_by_mitra.iterrows():
        mitra = row["Mitra"]
        site_count = row["Site ID Operator"]
        candidate_submit_count = row["Candidate NY Submit"]
        need_validation_count = row["Need Validation By Operator"]
        percentage = row["Percentage"]

        # Display mitra details and progress bar
        st.markdown(
            f'<div style="background-color: #FFA500; padding: 10px; border-radius: 10px; margin-bottom: 10px;">'
            f'<p style="font-size: 18px; text-align: center; color: #000000; margin: 0; font-weight: bold; text-shadow: 1px 1px 2px rgba(0,0,0,0.3);">{mitra}</p>'
            f'<p style="font-size: 14px; text-align: center; color: black; margin: 0;">Jumlah Site: {site_count}</p>'
            f'<p style="font-size: 14px; text-align: center; color: black; margin: 0;">Candidate NY Submit: {candidate_submit_count}</p>'
            f'<p style="font-size: 14px; text-align: center; color: black; margin: 0;">Need Validation By Operator: {need_validation_count}</p>'
            f'</div>',
            unsafe_allow_html=True
        )
        
        # Display progress bar with percentage at the end
        st.progress(int(percentage))
        st.markdown(f'<div style="text-align: right; font-size: 14px; color: black;">{int(percentage)}%</div>', unsafe_allow_html=True)

        # Menampilkan tabel untuk setiap Mitra
        mitra_df = df[df["Mitra"] == mitra][["Site ID Operator", "Kandidat P", "Kandidat Q", "Kandidat R", "Kandidat S", "LT STIP Release - Today", "LT Candidate Submitted NY Validation","Lead Time Validation"]]
        st.dataframe(mitra_df.reset_index(drop=True))

#fungsi menu generate TSSR ===================================================================
# Fungsi untuk membaca data dari file Excel
def load_excel(file):
    df = pd.read_excel(file)
    return df

# Fungsi untuk menggantikan placeholder di dalam teks
def replace_placeholders_in_text(text, data_row):
    for key, value in data_row.items():
        placeholder = "{{" + key + "}}"
        text = text.replace(placeholder, str(value))
    return text

# Fungsi untuk membuat dokumen Word berdasarkan template
def create_word_document(template_url, data_row):
    doc = Document(template_url)
    
    # Ganti placeholder di luar tabel
    for paragraph in doc.paragraphs:
        paragraph.text = replace_placeholders_in_text(paragraph.text, data_row)
    
    # Ganti placeholder di dalam tabel
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.text = replace_placeholders_in_text(paragraph.text, data_row)
    
    return doc

# Fungsi untuk mendapatkan data berdasarkan Site ID Operator
def get_site_data(sheet, site_id_operator):
    data = sheet.get_all_values()
    headers = data[1]
    rows = data[2:]
    
    df = pd.DataFrame(rows, columns=headers)
    site_data = df[df['Site ID Operator'].str.lower() == site_id_operator.lower()]
    
    if not site_data.empty:
        return site_data.iloc[0].to_dict()
    else:
        return None
    
def check_login(username,password):
    print("Checking login...")
    return username == "sitac" and password == "1234"

# Main program
def main():
    if 'login_state' not in st.session_state:
        st.session_state.login_state = False

    if not st.session_state.login_state:

        logotbg_url = 'https://raw.githubusercontent.com/samuelcode109/Presitac-Improvement-sipandai/main/tbg2.png'
        logoorang_url = 'https://raw.githubusercontent.com/samuelcode109/Presitac-Improvement-sipandai/main/logo2.png'
        tbglogo = Image.open(BytesIO(requests.get(logotbg_url).content))
        logoorang = Image.open(BytesIO(requests.get(logoorang_url).content))

        kolom3,kolom4=st.columns([1,3])
        with kolom3:
            st.image(logoorang, width=200)
            
        with kolom4:
            kolom1,kolom2=st.columns([2,2])
            with kolom1:
                st.title("Login:lock:")  # You can adjust the font_size here.
                st.title(":blue[SIPANDAI]")
            st.subheader(":blue[S]istem :blue[I]ntegrasi :blue[P]emantauan, :blue[AN]alisis dan :blue[D]okumen :blue[A]utomated Pre-S:blue[I]TAC 🚀" ) 
            with kolom2:
                st.image(tbglogo, width=200)

        st.info('Silahkan masukkan :red[Username] dan :red[Password] ', icon="📌")
        # Login section
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        login_button = st.button("Login")
    
        if login_button:
            if check_login(username, password):
                st.success("Login Successful!")
                st.session_state.login_state = True
                st.experimental_rerun()
            else:
                st.error("Invalid username or password. Please try again.")
    else:
        with st.sidebar:
             # Load the image/logo
            logo_url = 'https://raw.githubusercontent.com/samuelcode109/Presitac-Improvement-sipandai/main/tbg.png'
            logo_image = Image.open(BytesIO(requests.get(logo_url).content))

            # Display the image/logo in the sidebar
            st.sidebar.image(logo_image)
            
            menu = option_menu(
                menu_title="Main Menu",  # required
                options=["Progress Monitoring","Show TSSR","Input Data","Generate Doc TSSR"],  # required
                icons=["bar-chart-line","search","graph-up", "pencil"],  # required
                menu_icon="cast",  # optional
                default_index=0,  # optional
            )

        if menu == "Show TSSR":
            site_id_operator = st.text_input("Enter Site ID Operator")
            # Menghubungkan ke Google Sheets
            sheet = connect_to_google_sheets(spreadsheet_name, sheet_name="TSSR")
            if st.button("Search"):
                show_tssr_data(sheet, site_id_operator)
            
            if "tssr_data" in st.session_state and st.session_state["tssr_data"]:
                data = st.session_state["tssr_data"]
                st.write({
                    "Sonumb": data["Sonumb"],
                    "Site Name TBG": data["Site Name TBG"],
                    "Long NOM": data["Long NOM"],
                    "Lat NOM": data["Lat NOM"],
                    "Mitra": data["Mitra"]
                })
                
                # Input untuk radius SAR
                sar_radius = st.number_input("Enter SAR radius (meters)", min_value=0, value=500)

                # Tampilkan peta dengan lingkaran radius SAR dan titik tambahan
                lat_kand_p = float(data["Lat Kand P"]) if data["Lat Kand P"] else None
                long_kand_p = float(data["Long Kand P"]) if data["Long Kand P"] else None
                lat_kand_q = float(data["Lat Kand Q"]) if data["Lat Kand Q"] else None
                long_kand_q = float(data["Long Kand Q"]) if data["Long Kand Q"] else None
                lat_kand_r = float(data["Lat Kand R"]) if data["Lat Kand R"] else None
                long_kand_r = float(data["Long Kand R"]) if data["Long Kand R"] else None
                lat_kand_s = float(data["Lat Kand S"]) if data["Lat Kand S"] else None
                long_kand_s = float(data["Long Kand S"]) if data["Long Kand S"] else None

                show_map(
                    float(data["Lat NOM"]), 
                    float(data["Long NOM"]), 
                    lat_kand_p, long_kand_p, 
                    lat_kand_q, long_kand_q, 
                    lat_kand_r, long_kand_r, 
                    lat_kand_s, long_kand_s, 
                    radius=sar_radius
                )
                
                if st.button("Kandidat P"):
                    st.write({
                        "Alamat Site Kand P": data["Alamat Site Kand P"],
                        "Long Kand P": data["Long Kand P"],
                        "Lat Kand P": data["Lat Kand P"],
                        "Jarak kand P dari titik NOM (meter)": data["Jarak kand P dari titik NOM (meter)"],
                        "Kepemilikan Lahan Kand P": data["Kepemilikan Lahan Kand P"],
                        "Luas Lahan Kand P (meter)": data["Luas Lahan Kand P (meter)"],
                        "Nama Pemilik Lahan Kand P": data["Nama Pemilik Lahan Kand P"],
                        "Alamat Pemilik Lahan Kand P": data["Alamat Pemilik Lahan Kand P"],
                        "POI/Kawasan ramai disekitar Kand P": data["POI/Kawasan ramai disekitar Kand P"],
                        "Tipe Area sekitar Kand P": data["Tipe Area sekitar Kand P"],
                        "Tipe Site Kand P": data["Tipe Site Kand P"],
                        "Akses Menuju Kand P": data["Akses Menuju Kand P"],
                        "ukuran akses ke Kand P (meter)": data["ukuran akses ke Kand P (meter)"],
                        "Site terdekat dengan kandidat P (operator)": data["Site terdekat dengan kandidat P (operator)"],
                        "Jarak site terdekat dengan kandidat P (meter)": data["Jarak site terdekat dengan kandidat P (meter)"],
                        "Tipe site terdekat dengan kandidat P": data["Tipe site terdekat dengan kandidat P"],
                        "Tipe Tower terdekat Kand P": data["Tipe Tower terdekat Kand P"],
                        "Long site terdekat kand P": data["Long site terdekat kand P"],
                        "Lat site terdekat kand P": data["Lat site terdekat kand P"],
                        "Informasi lain Kand P": data["Informasi lain Kand P"],
                        "Jarak terdekat tiang PLN ke Kand P (meter)": data["Jarak terdekat tiang PLN ke Kand P (meter)"],
                        "Jarak terdekat transformator ke kand P (meter)": data["Jarak terdekat transformator ke kand P (meter)"],
                        "Validation Kand P": data["Validation Kand P"]
                    })
                if st.button("Kandidat Q"):
                    st.write({
                        "Alamat Site Kand Q": data["Alamat Site Kand Q"],
                        "Long Kand Q": data["Long Kand Q"],
                        "Lat Kand Q": data["Lat Kand Q"],
                        "Jarak kand Q dari titik NOM (meter)": data["Jarak kand Q dari titik NOM (meter)"],
                        "Kepemilikan Lahan Kand Q": data["Kepemilikan Lahan Kand Q"],
                        "Luas Lahan Kand Q (meter)": data["Luas Lahan Kand Q (meter)"],
                        "Nama Pemilik Lahan Kand Q": data["Nama Pemilik Lahan Kand Q"],
                        "Alamat Pemilik Lahan Kand Q": data["Alamat Pemilik Lahan Kand Q"],
                        "POI/Kawasan ramai disekitar Kand Q": data["POI/Kawasan ramai disekitar Kand Q"],
                        "Tipe Area sekitar Kand Q": data["Tipe Area sekitar Kand Q"],
                        "Tipe Site Kand Q": data["Tipe Site Kand Q"],
                        "Akses Menuju Kand Q": data["Akses Menuju Kand Q"],
                        "ukuran akses ke Kand Q (meter)": data["ukuran akses ke Kand Q (meter)"],
                        "Site terdekat dengan kandidat Q (operator)": data["Site terdekat dengan kandidat Q (operator)"],
                        "Jarak site terdekat dengan kandidat Q (meter)": data["Jarak site terdekat dengan kandidat Q (meter)"],
                        "Tipe site terdekat dengan kandidat Q": data["Tipe site terdekat dengan kandidat Q"],
                        "Tipe Tower terdekat Kand Q": data["Tipe Tower terdekat Kand Q"],
                        "Long site terdekat kand Q": data["Long site terdekat kand Q"],
                        "Lat site terdekat kand Q": data["Lat site terdekat kand Q"],
                        "Informasi lain Kand Q": data["Informasi lain Kand Q"],
                        "Jarak terdekat tiang PLN ke Kand Q (meter)": data["Jarak terdekat tiang PLN ke Kand Q (meter)"],
                        "Jarak terdekat transformator ke kand Q (meter)": data["Jarak terdekat transformator ke kand Q (meter)"],
                        "Validation Kand Q": data["Validation Kand Q"]
                    })

                if st.button("Kandidat R"):
                    st.write({
                        "Alamat Site Kand R": data["Alamat Site Kand R"],
                        "Long Kand R": data["Long Kand R"],
                        "Lat Kand R": data["Lat Kand R"],
                        "Jarak kand R dari titik NOM (meter)": data["Jarak kand R dari titik NOM (meter)"],
                        "Kepemilikan Lahan Kand R": data["Kepemilikan Lahan Kand R"],
                        "Luas Lahan Kand R (meter)": data["Luas Lahan Kand R (meter)"],
                        "Nama Pemilik Lahan Kand R": data["Nama Pemilik Lahan Kand R"],
                        "Alamat Pemilik Lahan Kand R": data["Alamat Pemilik Lahan Kand R"],
                        "POI/Kawasan ramai disekitar Kand R": data["POI/Kawasan ramai disekitar Kand R"],
                        "Tipe Area sekitar Kand R": data["Tipe Area sekitar Kand R"],
                        "Tipe Site Kand R": data["Tipe Site Kand R"],
                        "Akses Menuju Kand R": data["Akses Menuju Kand R"],
                        "ukuran akses ke Kand R (meter)": data["ukuran akses ke Kand R (meter)"],
                        "Site terdekat dengan kandidat R (operator)": data["Site terdekat dengan kandidat R (operator)"],
                        "Jarak site terdekat dengan kandidat R (meter)": data["Jarak site terdekat dengan kandidat R (meter)"],
                        "Tipe site terdekat dengan kandidat R": data["Tipe site terdekat dengan kandidat R"],
                        "Tipe Tower terdekat Kand R": data["Tipe Tower terdekat Kand R"],
                        "Long site terdekat kand R": data["Long site terdekat kand R"],
                        "Lat site terdekat kand R": data["Lat site terdekat kand R"],
                        "Informasi lain Kand R": data["Informasi lain Kand R"],
                        "Jarak terdekat tiang PLN ke Kand R (meter)": data["Jarak terdekat tiang PLN ke Kand R (meter)"],
                        "Jarak terdekat transformator ke kand R (meter)": data["Jarak terdekat transformator ke kand R (meter)"],
                        "Validation Kand R": data["Validation Kand R"]
                
                    })
                    
                    
                if st.button("Kandidat S"):
                    st.write({
                        "Alamat Site Kand S": data["Alamat Site Kand S"],
                        "Long Kand S": data["Long Kand S"],
                        "Lat Kand S": data["Lat Kand S"],
                        "Jarak kand S dari titik NOM (meter)": data["Jarak kand S dari titik NOM (meter)"],
                        "Kepemilikan Lahan Kand S": data["Kepemilikan Lahan Kand S"],
                        "Luas Lahan Kand S (meter)": data["Luas Lahan Kand S (meter)"],
                        "Nama Pemilik Lahan Kand S": data["Nama Pemilik Lahan Kand S"],
                        "Alamat Pemilik Lahan Kand S": data["Alamat Pemilik Lahan Kand S"],
                        "POI/Kawasan ramai disekitar Kand S": data["POI/Kawasan ramai disekitar Kand S"],
                        "Tipe Area sekitar Kand S": data["Tipe Area sekitar Kand S"],
                        "Tipe Site Kand S": data["Tipe Site Kand S"],
                        "Akses Menuju Kand S": data["Akses Menuju Kand S"],
                        "ukuran akses ke Kand S (meter)": data["ukuran akses ke Kand S (meter)"],
                        "Site terdekat dengan kandidat S (operator)": data["Site terdekat dengan kandidat S (operator)"],
                        "Jarak site terdekat dengan kandidat S (meter)": data["Jarak site terdekat dengan kandidat S (meter)"],
                        "Tipe site terdekat dengan kandidat S": data["Tipe site terdekat dengan kandidat S"],
                        "Tipe Tower terdekat Kand S": data["Tipe Tower terdekat Kand S"],
                        "Long site terdekat kand S": data["Long site terdekat kand S"],
                        "Lat site terdekat kand S": data["Lat site terdekat kand S"],
                        "Informasi lain Kand S": data["Informasi lain Kand S"],
                        "Jarak terdekat tiang PLN ke Kand S (meter)": data["Jarak terdekat tiang PLN ke Kand S (meter)"],
                        "Jarak terdekat transformator ke kand S (meter)": data["Jarak terdekat transformator ke kand S (meter)"],
                        "Validation Kand S": data["Validation Kand S"]
                    })

            elif "tssr_data" in st.session_state:
                st.write("Site ID Operator not found.")
                
        elif menu == "Progress Monitoring":
            # Menghubungkan ke sheet "Progress"
            sheet = connect_to_google_sheets(spreadsheet_name, sheet_name="Progress")
            show_progress_dashboard(sheet)
            
        elif menu == "Input Data":
            st.title("Input Data")
            
            # Menghubungkan ke Google Sheets
            sheet = connect_to_google_sheets(spreadsheet_name, sheet_name="TSSR")
            
            site_id_operator = st.text_input("Enter Site ID Operator")
            
            if site_id_operator:
                site_data = get_site_data(sheet, site_id_operator)
                
                if site_data:
                    st.write("Sonumb:", site_data["Sonumb"])
                    st.write("Site Name TBG:", site_data["Site Name TBG"])
                    st.write("Long NOM:", site_data["Long NOM"])
                    st.write("Lat NOM:", site_data["Lat NOM"])
                    st.write("Mitra:", site_data["Mitra"])
                    
                    st.write("Input Data for Kandidat P")
                    alamat_kand_p = st.text_input("Alamat Site Kand P", value=site_data.get("Alamat Site Kand P", ""))
                    long_kand_p = st.text_input("Long Kand P", value=site_data.get("Long Kand P", ""))
                    lat_kand_p = st.text_input("Lat Kand P", value=site_data.get("Lat Kand P", ""))
                    p4=st.text_input("Jarak kand P dari titik NOM (meter)", value=site_data.get("Jarak kand P dari titik NOM (meter)", ""))
                    p5=st.text_input("Kepemilikan Lahan Kand P", value=site_data.get("Kepemilikan Lahan Kand P", ""))
                    p6=st.text_input("Luas Lahan Kand P (meter)", value=site_data.get("Luas Lahan Kand P (meter)", ""))
                    p7=st.text_input("Nama Pemilik Lahan Kand P", value=site_data.get("Nama Pemilik Lahan Kand P", ""))
                    p8=st.text_input("Alamat Pemilik Lahan Kand P", value=site_data.get("Alamat Pemilik Lahan Kand P", ""))
                    p9=st.text_input("POI/Kawasan ramai disekitar Kand P", value=site_data.get("POI/Kawasan ramai disekitar Kand P", ""))
                    p10=st.text_input("Tipe Area sekitar Kand P", value=site_data.get("Tipe Area sekitar Kand P", ""))
                    p11=st.text_input("Tipe Site Kand P", value=site_data.get("Tipe Site Kand P", ""))
                    p12=st.text_input("Akses Menuju Kand P", value=site_data.get("Akses Menuju Kand P", ""))
                    p13=st.text_input("ukuran akses ke Kand P (meter)", value=site_data.get("ukuran akses ke Kand P (meter)", ""))
                    p14=st.text_input("Site terdekat dengan kandidat P (operator)", value=site_data.get("Site terdekat dengan kandidat P (operator)", ""))
                    p15=st.text_input("Jarak site terdekat dengan kandidat P (meter)", value=site_data.get("Jarak site terdekat dengan kandidat P (meter)", ""))
                    p16=st.text_input("Tipe site terdekat dengan kandidat P", value=site_data.get("Tipe site terdekat dengan kandidat P", ""))
                    p17=st.text_input("Tipe Tower terdekat Kand P", value=site_data.get("Tipe Tower terdekat Kand P", ""))
                    p18=st.text_input("Long site terdekat kand P", value=site_data.get("Long site terdekat kand P", ""))
                    p19=st.text_input("Lat site terdekat kand P", value=site_data.get("Lat site terdekat kand P", ""))
                    p20=st.text_input("Informasi lain Kand P", value=site_data.get("Informasi lain Kand P", ""))
                    p21=st.text_input("Jarak terdekat tiang PLN ke Kand P (meter)", value=site_data.get("Jarak terdekat tiang PLN ke Kand P (meter)", ""))
                    p22=st.text_input("Jarak terdekat transformator ke kand P (meter)", value=site_data.get("Jarak terdekat transformator ke kand P (meter)", ""))
                    
                    st.write("Input Data for Kandidat Q")
                    alamat_kand_q = st.text_input("Alamat Site Kand Q", value=site_data.get("Alamat Site Kand Q", ""))
                    long_kand_q = st.text_input("Long Kand Q", value=site_data.get("Long Kand Q", ""))
                    lat_kand_q = st.text_input("Lat Kand Q", value=site_data.get("Lat Kand Q", ""))
                    q4=st.text_input("Jarak kand Q dari titik NOM (meter)", value=site_data.get("Jarak kand Q dari titik NOM (meter)", ""))
                    q5=st.text_input("Kepemilikan Lahan Kand Q", value=site_data.get("Kepemilikan Lahan Kand Q", ""))
                    q6=st.text_input("Luas Lahan Kand Q (meter)", value=site_data.get("Luas Lahan Kand Q (meter)", ""))
                    q7=st.text_input("Nama Pemilik Lahan Kand Q", value=site_data.get("Nama Pemilik Lahan Kand Q", ""))
                    q8=st.text_input("Alamat Pemilik Lahan Kand Q", value=site_data.get("Alamat Pemilik Lahan Kand Q", ""))
                    q9=st.text_input("POI/Kawasan ramai disekitar Kand Q", value=site_data.get("POI/Kawasan ramai disekitar Kand Q", ""))
                    q10=st.text_input("Tipe Area sekitar Kand Q", value=site_data.get("Tipe Area sekitar Kand Q", ""))
                    q11=st.text_input("Tipe Site Kand Q", value=site_data.get("Tipe Site Kand Q", ""))
                    q12=st.text_input("Akses Menuju Kand Q", value=site_data.get("Akses Menuju Kand Q", ""))
                    q13=st.text_input("ukuran akses ke Kand Q (meter)", value=site_data.get("ukuran akses ke Kand Q (meter)", ""))
                    q14=st.text_input("Site terdekat dengan kandidat Q (operator)", value=site_data.get("Site terdekat dengan kandidat Q (operator)", ""))
                    q15=st.text_input("Jarak site terdekat dengan kandidat Q (meter)", value=site_data.get("Jarak site terdekat dengan kandidat Q (meter)", ""))
                    q16=st.text_input("Tipe site terdekat dengan kandidat Q", value=site_data.get("Tipe site terdekat dengan kandidat Q", ""))
                    q17=st.text_input("Tipe Tower terdekat Kand Q", value=site_data.get("Tipe Tower terdekat Kand Q", ""))
                    q18=st.text_input("Long site terdekat kand Q", value=site_data.get("Long site terdekat kand Q", ""))
                    q19=st.text_input("Lat site terdekat kand Q", value=site_data.get("Lat site terdekat kand Q", ""))
                    q20=st.text_input("Informasi lain Kand Q", value=site_data.get("Informasi lain Kand Q", ""))
                    q21=st.text_input("Jarak terdekat tiang PLN ke Kand Q (meter)", value=site_data.get("Jarak terdekat tiang PLN ke Kand Q (meter)", ""))
                    q22=st.text_input("Jarak terdekat transformator ke kand Q (meter)", value=site_data.get("Jarak terdekat transformator ke kand Q (meter)", ""))
                    
                    st.write("Input Data for Kandidat R")
                    alamat_kand_r = st.text_input("Alamat Site Kand R", value=site_data.get("Alamat Site Kand R", ""))
                    long_kand_r = st.text_input("Long Kand R", value=site_data.get("Long Kand R", ""))
                    lat_kand_r = st.text_input("Lat Kand R", value=site_data.get("Lat Kand R", ""))
                    r4=st.text_input("Jarak kand R dari titik NOM (meter)", value=site_data.get("Jarak kand R dari titik NOM (meter)", ""))
                    r5=st.text_input("Kepemilikan Lahan Kand R", value=site_data.get("Kepemilikan Lahan Kand R", ""))
                    r6=st.text_input("Luas Lahan Kand R (meter)", value=site_data.get("Luas Lahan Kand R (meter)", ""))
                    r7=st.text_input("Nama Pemilik Lahan Kand R", value=site_data.get("Nama Pemilik Lahan Kand R", ""))
                    r8=st.text_input("Alamat Pemilik Lahan Kand R", value=site_data.get("Alamat Pemilik Lahan Kand R", ""))
                    r9=st.text_input("POI/Kawasan ramai disekitar Kand R", value=site_data.get("POI/Kawasan ramai disekitar Kand R", ""))
                    r10=st.text_input("Tipe Area sekitar Kand R", value=site_data.get("Tipe Area sekitar Kand R", ""))
                    r11=st.text_input("Tipe Site Kand R", value=site_data.get("Tipe Site Kand R", ""))
                    r12=st.text_input("Akses Menuju Kand R", value=site_data.get("Akses Menuju Kand R", ""))
                    r13=st.text_input("ukuran akses ke Kand R (meter)", value=site_data.get("ukuran akses ke Kand R (meter)", ""))
                    r14=st.text_input("Site terdekat dengan kandidat R (operator)", value=site_data.get("Site terdekat dengan kandidat R (operator)", ""))
                    r15=st.text_input("Jarak site terdekat dengan kandidat R (meter)", value=site_data.get("Jarak site terdekat dengan kandidat R (meter)", ""))
                    r16=st.text_input("Tipe site terdekat dengan kandidat R", value=site_data.get("Tipe site terdekat dengan kandidat R", ""))
                    r17=st.text_input("Tipe Tower terdekat Kand R", value=site_data.get("Tipe Tower terdekat Kand R", ""))
                    r18=st.text_input("Long site terdekat kand R", value=site_data.get("Long site terdekat kand R", ""))
                    r19=st.text_input("Lat site terdekat kand R", value=site_data.get("Lat site terdekat kand R", ""))
                    r20=st.text_input("Informasi lain Kand R", value=site_data.get("Informasi lain Kand R", ""))
                    r21=st.text_input("Jarak terdekat tiang PLN ke Kand R (meter)", value=site_data.get("Jarak terdekat tiang PLN ke Kand R (meter)", ""))
                    r22=st.text_input("Jarak terdekat transformator ke kand R (meter)", value=site_data.get("Jarak terdekat transformator ke kand R (meter)", ""))
                    
                    if st.button("Submit"):
                        # Update data di sheet
                        headers = sheet.row_values(2)
                        alamat_kand_p_col = headers.index("Alamat Site Kand P") + 1
                        long_kand_p_col = headers.index("Long Kand P") + 1
                        lat_kand_p_col = headers.index("Lat Kand P") + 1
                        p4_col = headers.index("Jarak kand P dari titik NOM (meter)") + 1
                        p5_col = headers.index("Kepemilikan Lahan Kand P") + 1
                        p6_col = headers.index("Luas Lahan Kand P (meter)") + 1
                        p7_col = headers.index("Nama Pemilik Lahan Kand P") + 1
                        p8_col = headers.index("Alamat Pemilik Lahan Kand P") + 1
                        p9_col = headers.index("POI/Kawasan ramai disekitar Kand P") + 1
                        p10_col = headers.index("Tipe Area sekitar Kand P") + 1
                        p11_col = headers.index("Tipe Site Kand P") + 1
                        p12_col = headers.index("Akses Menuju Kand P") + 1
                        p13_col = headers.index("ukuran akses ke Kand P (meter)") + 1
                        p14_col = headers.index("Site terdekat dengan kandidat P (operator)") + 1
                        p15_col = headers.index("Jarak site terdekat dengan kandidat P (meter)") + 1
                        p16_col = headers.index("Tipe site terdekat dengan kandidat P") + 1
                        p17_col = headers.index("Tipe Tower terdekat Kand P") + 1
                        p18_col = headers.index("Long site terdekat kand P") + 1
                        p19_col = headers.index("Lat site terdekat kand P") + 1
                        p20_col = headers.index("Informasi lain Kand P") + 1
                        p21_col = headers.index("Jarak terdekat tiang PLN ke Kand P (meter)") + 1
                        p22_col = headers.index("Jarak terdekat transformator ke kand P (meter)") + 1
                        
                        alamat_kand_q_col = headers.index("Alamat Site Kand Q") + 1
                        long_kand_q_col = headers.index("Long Kand Q") + 1
                        lat_kand_q_col = headers.index("Lat Kand Q") + 1
                        q4_col = headers.index("Jarak kand Q dari titik NOM (meter)") + 1
                        q5_col = headers.index("Kepemilikan Lahan Kand Q") + 1
                        q6_col = headers.index("Luas Lahan Kand Q (meter)") + 1
                        q7_col = headers.index("Nama Pemilik Lahan Kand Q") + 1
                        q8_col = headers.index("Alamat Pemilik Lahan Kand Q") + 1
                        q9_col = headers.index("POI/Kawasan ramai disekitar Kand Q") + 1
                        q10_col = headers.index("Tipe Area sekitar Kand Q") + 1
                        q11_col = headers.index("Tipe Site Kand Q") + 1
                        q12_col = headers.index("Akses Menuju Kand Q") + 1
                        q13_col = headers.index("ukuran akses ke Kand Q (meter)") + 1
                        q14_col = headers.index("Site terdekat dengan kandidat Q (operator)") + 1
                        q15_col = headers.index("Jarak site terdekat dengan kandidat Q (meter)") + 1
                        q16_col = headers.index("Tipe site terdekat dengan kandidat Q") + 1
                        q17_col = headers.index("Tipe Tower terdekat Kand Q") + 1
                        q18_col = headers.index("Long site terdekat kand Q") + 1
                        q19_col = headers.index("Lat site terdekat kand Q") + 1
                        q20_col = headers.index("Informasi lain Kand Q") + 1
                        q21_col = headers.index("Jarak terdekat tiang PLN ke Kand Q (meter)") + 1
                        q22_col = headers.index("Jarak terdekat transformator ke kand Q (meter)") + 1
                        
                        alamat_kand_r_col = headers.index("Alamat Site Kand R") + 1
                        long_kand_r_col = headers.index("Long Kand R") + 1
                        lat_kand_r_col = headers.index("Lat Kand R") + 1
                        r4_col = headers.index("Jarak kand R dari titik NOM (meter)") + 1
                        r5_col = headers.index("Kepemilikan Lahan Kand R") + 1
                        r6_col = headers.index("Luas Lahan Kand R (meter)") + 1
                        r7_col = headers.index("Nama Pemilik Lahan Kand R") + 1
                        r8_col = headers.index("Alamat Pemilik Lahan Kand R") + 1
                        r9_col = headers.index("POI/Kawasan ramai disekitar Kand R") + 1
                        r10_col = headers.index("Tipe Area sekitar Kand R") + 1
                        r11_col = headers.index("Tipe Site Kand R") + 1
                        r12_col = headers.index("Akses Menuju Kand R") + 1
                        r13_col = headers.index("ukuran akses ke Kand R (meter)") + 1
                        r14_col = headers.index("Site terdekat dengan kandidat R (operator)") + 1
                        r15_col = headers.index("Jarak site terdekat dengan kandidat R (meter)") + 1
                        r16_col = headers.index("Tipe site terdekat dengan kandidat R") + 1
                        r17_col = headers.index("Tipe Tower terdekat Kand R") + 1
                        r18_col = headers.index("Long site terdekat kand R") + 1
                        r19_col = headers.index("Lat site terdekat kand R") + 1
                        r20_col = headers.index("Informasi lain Kand R") + 1
                        r21_col = headers.index("Jarak terdekat tiang PLN ke Kand R (meter)") + 1
                        r22_col = headers.index("Jarak terdekat transformator ke kand R (meter)") + 1
                        
                        cell = sheet.find(site_id_operator)
                        if cell:
                            row = cell.row
                            sheet.update_cell(row, alamat_kand_p_col, alamat_kand_p)
                            sheet.update_cell(row, long_kand_p_col, long_kand_p)
                            sheet.update_cell(row, lat_kand_p_col, lat_kand_p)
                            sheet.update_cell(row, p4_col, p4)
                            sheet.update_cell(row, p5_col, p5)
                            sheet.update_cell(row, p6_col, p6)
                            sheet.update_cell(row, p7_col, p7)
                            sheet.update_cell(row, p8_col, p8)
                            sheet.update_cell(row, p9_col, p9)
                            sheet.update_cell(row, p10_col, p10)
                            sheet.update_cell(row, p11_col, p11)
                            sheet.update_cell(row, p12_col, p12)
                            sheet.update_cell(row, p13_col, p13)
                            sheet.update_cell(row, p14_col, p14)
                            sheet.update_cell(row, p15_col, p15)
                            sheet.update_cell(row, p16_col, p16)
                            sheet.update_cell(row, p17_col, p17)
                            sheet.update_cell(row, p18_col, p18)
                            sheet.update_cell(row, p19_col, p19)
                            sheet.update_cell(row, p20_col, p20)
                            sheet.update_cell(row, p21_col, p21)
                            sheet.update_cell(row, p22_col, p22)
                            
                            sheet.update_cell(row, alamat_kand_q_col, alamat_kand_q)
                            sheet.update_cell(row, long_kand_q_col, long_kand_q)
                            sheet.update_cell(row, lat_kand_q_col, lat_kand_q)
                            sheet.update_cell(row, q4_col, q4)
                            sheet.update_cell(row, q5_col, q5)
                            sheet.update_cell(row, q6_col, q6)
                            sheet.update_cell(row, q7_col, q7)
                            sheet.update_cell(row, q8_col, q8)
                            sheet.update_cell(row, q9_col, q9)
                            sheet.update_cell(row, q10_col, q10)
                            sheet.update_cell(row, q11_col, q11)
                            sheet.update_cell(row, q12_col, q12)
                            sheet.update_cell(row, q13_col, q13)
                            sheet.update_cell(row, q14_col, q14)
                            sheet.update_cell(row, q15_col, q15)
                            sheet.update_cell(row, q16_col, q16)
                            sheet.update_cell(row, q17_col, q17)
                            sheet.update_cell(row, q18_col, q18)
                            sheet.update_cell(row, q19_col, q19)
                            sheet.update_cell(row, q20_col, q20)
                            sheet.update_cell(row, q21_col, q21)
                            sheet.update_cell(row, q22_col, q22)
                            
                            sheet.update_cell(row, alamat_kand_r_col, alamat_kand_r)
                            sheet.update_cell(row, long_kand_r_col, long_kand_r)
                            sheet.update_cell(row, lat_kand_r_col, lat_kand_r)
                            sheet.update_cell(row, r4_col, r4)
                            sheet.update_cell(row, r5_col, r5)
                            sheet.update_cell(row, r6_col, r6)
                            sheet.update_cell(row, r7_col, r7)
                            sheet.update_cell(row, r8_col, r8)
                            sheet.update_cell(row, r9_col, r9)
                            sheet.update_cell(row, r10_col, r10)
                            sheet.update_cell(row, r11_col, r11)
                            sheet.update_cell(row, r12_col, r12)
                            sheet.update_cell(row, r13_col, r13)
                            sheet.update_cell(row, r14_col, r14)
                            sheet.update_cell(row, r15_col, r15)
                            sheet.update_cell(row, r16_col, r16)
                            sheet.update_cell(row, r17_col, r17)
                            sheet.update_cell(row, r18_col, r18)
                            sheet.update_cell(row, r19_col, r19)
                            sheet.update_cell(row, r20_col, r20)
                            sheet.update_cell(row, r21_col, r21)
                            sheet.update_cell(row, r22_col, r22)

                            st.success("Data successfully submitted!")
                else:
                    st.error("Site ID Operator not found!")
                    
                    
        elif menu == "Generate Doc TSSR":
            st.title("Auto Generate Doc TSSR")

            # Upload file Excel
            uploaded_file = st.file_uploader("Choose an Excel file", type="xlsx")
            template_url = 'https://raw.githubusercontent.com/samuelcode109/Presitac-Improvement-sipandai/main/tssr1.docx' # URL ke template Word Anda


            if uploaded_file is not None:
                df = load_excel(uploaded_file)
                st.write("Excel data:", df)
                
                if st.button("Create Word Documents"):
                    # Inisialisasi buffer untuk ZIP file
                    zip_buffer = BytesIO()
                    
                    with zipfile.ZipFile(zip_buffer, 'a', zipfile.ZIP_DEFLATED) as zip_file:
                        for idx, row in df.iterrows():
                            # Buat dokumen Word untuk setiap baris
                            doc = create_word_document(template_url, row)
                            doc_buffer = BytesIO()
                            doc.save(doc_buffer)
                            doc_buffer.seek(0)
                            
                            # Nama file untuk setiap dokumen
                            site_name = row['Site Name TBG']
                            doc_name = f"site_{site_name}.docx"
                            
                            # Tambahkan dokumen ke ZIP file
                            zip_file.writestr(doc_name, doc_buffer.getvalue())
                    
                    # Siapkan ZIP file untuk diunduh
                    zip_buffer.seek(0)
                    st.download_button(
                        label="Download Word Documents",
                        data=zip_buffer,
                        file_name="documents.zip",
                        mime="application/zip"
                    )

if __name__ == "__main__":
    main()